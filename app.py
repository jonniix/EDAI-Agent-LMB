<<<<<<< HEAD
import os
import time
import tempfile
import numpy as np
import streamlit as st
from pydub import AudioSegment
import whisper
from transformers import pipeline
from nltk.tokenize import sent_tokenize
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import torch
import ffmpeg
from typing import Optional, Tuple, List
from streamlit.components.v1 import html
import base64

# Set up NLTK
try:
    sent_tokenize("Test sentence.")
except:
    import nltk
    nltk.download('punkt')

# App configuration
st.set_page_config(
    page_title="🧠 AI Voice Transcriber & Reporter",
    page_icon="🎙️",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Constants
SUPPORTED_AUDIO_EXTENSIONS = [".mp3", ".wav", ".m4a", ".ogg", ".aac"]
SUPPORTED_VIDEO_EXTENSIONS = [".mp4", ".mov", ".avi"]
SUMMARIZATION_MODEL = "facebook/bart-large-cnn"
WHISPER_MODEL = "base"  # Default model
MAX_FILE_SIZE_MB = 100
CHUNK_SIZE_SECONDS = 20  # For live transcription

# Session state initialization
if 'transcription' not in st.session_state:
    st.session_state.transcription = None
if 'summary' not in st.session_state:
    st.session_state.summary = None
if 'translation' not in st.session_state:
    st.session_state.translation = None
if 'whisper_model' not in st.session_state:
    st.session_state.whisper_model = None
if 'summarizer' not in st.session_state:
    st.session_state.summarizer = None
if 'translation_pipeline' not in st.session_state:
    st.session_state.translation_pipeline = None
if 'transcription_chunks' not in st.session_state:
    st.session_state.transcription_chunks = []
if 'translation_chunks' not in st.session_state:
    st.session_state.translation_chunks = []
if 'processing_complete' not in st.session_state:
    st.session_state.processing_complete = False
if 'current_chunk' not in st.session_state:
    st.session_state.current_chunk = 0

# Helper functions
def convert_to_wav(input_file: str) -> str:
    """Convert any audio file to WAV format for processing"""
    output_file = os.path.join(tempfile.gettempdir(), f"converted_{int(time.time())}.wav")
    
    try:
        audio = AudioSegment.from_file(input_file)
        audio.export(output_file, format="wav")
    except Exception as e:
        st.error(f"Conversion failed with pydub: {str(e)}. Trying ffmpeg...")
        try:
            (
                ffmpeg
                .input(input_file)
                .output(output_file, ac=1, ar=16000)
                .run(quiet=True, overwrite_output=True)
            )
        except Exception as e:
            st.error(f"FFmpeg conversion also failed: {str(e)}")
            return None
    
    return output_file

def split_audio(file_path: str, chunk_size: int = CHUNK_SIZE_SECONDS) -> List[str]:
    """Split audio into chunks for live processing"""
    try:
        audio = AudioSegment.from_file(file_path)
        duration_ms = len(audio)
        chunk_ms = chunk_size * 1000
        chunks = []
        
        for i in range(0, duration_ms, chunk_ms):
            chunk = audio[i:i+chunk_ms]
            chunk_path = os.path.join(tempfile.gettempdir(), f"chunk_{i//1000}.wav")
            chunk.export(chunk_path, format="wav")
            chunks.append(chunk_path)
        
        return chunks
    except Exception as e:
        st.error(f"Failed to split audio: {str(e)}")
        return []

def load_whisper_model(model_name: str = "base") -> whisper.Whisper:
    """Load Whisper model with caching"""
    if st.session_state.whisper_model is None or st.session_state.whisper_model.name != model_name:
        with st.spinner(f"Loading Whisper {model_name} model (first time may take a while)..."):
            st.session_state.whisper_model = whisper.load_model(model_name)
    return st.session_state.whisper_model

def load_summarizer() -> pipeline:
    """Load summarization pipeline with caching"""
    if st.session_state.summarizer is None:
        with st.spinner("Loading summarization model..."):
            st.session_state.summarizer = pipeline("summarization", model=SUMMARIZATION_MODEL)
    return st.session_state.summarizer

def load_translation_pipeline(target_lang: str = "en") -> Optional[pipeline]:
    """Load translation pipeline for target language"""
    model_map = {
        "en": "Helsinki-NLP/opus-mt-mul-en",
        "it": "Helsinki-NLP/opus-mt-en-it",
        "fr": "Helsinki-NLP/opus-mt-en-fr",
        "es": "Helsinki-NLP/opus-mt-en-es",
        "de": "Helsinki-NLP/opus-mt-en-de"
    }
    
    if target_lang not in model_map:
        return None
    
    try:
        if st.session_state.translation_pipeline is None or st.session_state.translation_pipeline.model.config.lang2 != target_lang:
            with st.spinner(f"Loading translation model for {target_lang}..."):
                st.session_state.translation_pipeline = pipeline(
                    "translation", 
                    model=model_map[target_lang]
                )
        return st.session_state.translation_pipeline
    except Exception as e:
        st.error(f"Failed to load translation model: {str(e)}")
        return None

def transcribe_chunk(chunk_path: str, model: whisper.Whisper) -> str:
    """Transcribe a single audio chunk"""
    try:
        result = model.transcribe(chunk_path)
        return result["text"]
    except Exception as e:
        st.error(f"Error transcribing chunk: {str(e)}")
        return ""

def summarize_text(text: str, summary_length: str = "medium") -> str:
    """Summarize text using BART model"""
    summarizer = load_summarizer()
    
    length_params = {
        "short": {"min_length": 30, "max_length": 80},
        "medium": {"min_length": 60, "max_length": 150},
        "long": {"min_length": 120, "max_length": 300}
    }
    
    params = length_params.get(summary_length, length_params["medium"])
    
    sentences = sent_tokenize(text)
    chunks = []
    current_chunk = ""
    
    for sentence in sentences:
        if len(current_chunk) + len(sentence) < 1000:
            current_chunk += " " + sentence
        else:
            chunks.append(current_chunk)
            current_chunk = sentence
    
    if current_chunk:
        chunks.append(current_chunk)
    
    summaries = []
    progress_bar = st.progress(0)
    for i, chunk in enumerate(chunks):
        summary = summarizer(chunk, **params, do_sample=False)
        summaries.append(summary[0]['summary_text'])
        progress_bar.progress((i + 1) / len(chunks))
    
    return " ".join(summaries)

def translate_text(text: str, target_lang: str) -> str:
    """Translate text to target language"""
    translator = load_translation_pipeline(target_lang)
    if not translator:
        return None
    
    sentences = sent_tokenize(text)
    chunks = []
    current_chunk = ""
    
    for sentence in sentences:
        if len(current_chunk) + len(sentence) < 1000:
            current_chunk += " " + sentence
        else:
            chunks.append(current_chunk)
            current_chunk = sentence
    
    if current_chunk:
        chunks.append(current_chunk)
    
    translations = []
    progress_bar = st.progress(0)
    for i, chunk in enumerate(chunks):
        translation = translator(chunk)
        translations.append(translation[0]['translation_text'])
        progress_bar.progress((i + 1) / len(chunks))
    
    return " ".join(translations)

def create_word_document(transcription: str, summary: str, translation: Optional[str] = None) -> Document:
    """Create a professional Word document with the results"""
    doc = Document()
    
    # Set document margins
    for section in doc.sections:
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
    
    # Add title
    title = doc.add_paragraph()
    title_run = title.add_run("AI Voice Transcription Report")
    title_run.bold = True
    title_run.font.size = Pt(16)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Add subtitle with date
    doc.add_paragraph(time.strftime("%B %d, %Y"), style='Intense Quote')
    doc.add_paragraph("\n")
    
    # Add summary section
    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph(summary)
    doc.add_paragraph("\n")
    
    # Add transcription section
    doc.add_heading("Full Transcription", level=1)
    doc.add_paragraph(transcription)
    
    # Add translation if available
    if translation:
        doc.add_paragraph("\n")
        doc.add_heading("Translation", level=1)
        doc.add_paragraph(translation)
    
    # Add footer
    doc.add_paragraph("\n")
    footer = doc.add_paragraph()
    footer_run = footer.add_run("Generated by AI Voice Transcriber & Reporter")
    footer_run.italic = True
    footer_run.font.size = Pt(10)
    footer.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    return doc

def get_file_duration(file_path: str) -> float:
    """Get duration of audio/video file in seconds"""
    try:
        probe = ffmpeg.probe(file_path)
        duration = float(probe['format']['duration'])
        return duration
    except:
        return 0

def animate_text(text: str, speed: int = 5) -> None:
    """Create typing animation effect for text"""
    placeholder = st.empty()
    displayed_text = ""
    for char in text:
        displayed_text += char
        placeholder.markdown(f"""
        <div style="
            background-color: rgba(0, 0, 0, 0.05);
            padding: 12px;
            border-radius: 8px;
            margin: 8px 0;
            line-height: 1.6;
        ">
            {displayed_text}
        </div>
        """, unsafe_allow_html=True)
        time.sleep(1/speed)
    return displayed_text

# UI Components
def setup_sidebar():
    """Configure the sidebar options"""
    with st.sidebar:
        st.title("⚙️ Settings")
        
        # Model selection
        st.subheader("Model Options")
        whisper_model = st.selectbox(
            "Whisper Model Size",
            options=["tiny", "base", "small", "medium", "large"],
            index=1,
            help="Larger models are more accurate but require more memory and time"
        )
        
        summary_length = st.selectbox(
            "Summary Detail Level",
            options=["Brief", "Medium", "Detailed"],
            index=1,
            help="Control the length and detail of the generated summary"
        )
        
        # Translation options
        st.subheader("🌍 Translation")
        translate = st.checkbox("Enable Translation", False)
        target_lang = st.selectbox(
            "Target Language",
            options=["🇬🇧 English", "🇮🇹 Italian", "🇫🇷 French", "🇪🇸 Spanish", "🇩🇪 German"],
            index=0,
            disabled=not translate
        )
        
        # Theme options
        st.subheader("🎨 UI Preferences")
        theme = st.selectbox(
            "Color Theme",
            options=["Auto (System)", "Light", "Dark"],
            index=0
        )
        
        # Processing options
        st.subheader("⚡ Processing")
        live_mode = st.checkbox(
            "Live Transcription Mode", 
            True,
            help="Show transcription as it's being generated"
        )
        
        st.markdown("---")
        st.markdown("""
        **About this app:**
        - Powered by OpenAI Whisper for transcription
        - Uses Facebook BART for summarization
        - Helsinki-NLP for translation
        - Works fully offline after initial setup
        """)
    
    return {
        "whisper_model": whisper_model.lower(),
        "summary_length": summary_length.lower(),
        "translate": translate,
        "target_lang": target_lang[3:5].lower() if translate else None,
        "live_mode": live_mode,
        "theme": theme
    }

def display_file_info(file_path: str):
    """Display metadata about the uploaded file"""
    file_name = os.path.basename(file_path)
    file_size = os.path.getsize(file_path) / (1024 * 1024)  # in MB
    duration = get_file_duration(file_path)
    
    st.subheader("📄 File Information")
    col1, col2, col3 = st.columns(3)
    with col1:
        st.metric("File Name", file_name)
    with col2:
        st.metric("File Size", f"{file_size:.2f} MB")
    with col3:
        st.metric("Duration", f"{duration:.2f} sec" if duration > 0 else "Unknown")

def display_live_transcription():
    """Display transcription chunks as they're processed"""
    st.subheader("🔴 Live Transcription Progress")
    progress_bar = st.progress(0)
    status_text = st.empty()
    
    total_chunks = len(st.session_state.transcription_chunks)
    model = load_whisper_model(st.session_state.settings["whisper_model"])
    
    for i in range(st.session_state.current_chunk, total_chunks):
        chunk_path = st.session_state.transcription_chunks[i]
        
        # Update progress
        progress = (i + 1) / total_chunks
        progress_bar.progress(progress)
        status_text.markdown(f"""
        <div style="background-color: rgba(0, 0, 0, 0.05); padding: 12px; border-radius: 8px; margin: 8px 0;">
            Processing chunk {i+1} of {total_chunks}...
        </div>
        """, unsafe_allow_html=True)
        
        # Transcribe chunk
        chunk_text = transcribe_chunk(chunk_path, model)
        
        # Display chunk with animation
        st.markdown(f"""
        <div style="
            background-color: rgba(0, 0, 0, 0.05);
            padding: 12px;
            border-radius: 8px;
            margin: 8px 0;
            border-left: 4px solid #4CAF50;
        ">
            <strong>Chunk {i+1}:</strong><br>
        </div>
        """, unsafe_allow_html=True)
        
        animate_text(chunk_text)
        
        # Update full transcription
        if st.session_state.transcription is None:
            st.session_state.transcription = ""
        st.session_state.transcription += " " + chunk_text
        st.session_state.current_chunk = i + 1
    
    # Complete processing
    progress_bar.empty()
    status_text.success("✅ Transcription complete!")
    st.session_state.processing_complete = True

def display_final_results():
    """Display the final transcription, summary, and translation"""
    st.subheader("📝 Final Results")
    
    # Summary section
    with st.expander("📋 Executive Summary", expanded=True):
        if st.session_state.summary is None:
            with st.spinner("Generating summary..."):
                st.session_state.summary = summarize_text(
                    st.session_state.transcription,
                    st.session_state.settings["summary_length"]
                )
        st.markdown(f"""
        <div style="
            background-color: rgba(0, 0, 0, 0.05);
            padding: 16px;
            border-radius: 8px;
            margin: 8px 0;
            line-height: 1.6;
        ">
            {st.session_state.summary}
        </div>
        """, unsafe_allow_html=True)
    
    # Transcription section
    with st.expander("🎧 Full Transcription", expanded=False):
        st.markdown(f"""
        <div style="
            background-color: rgba(0, 0, 0, 0.05);
            padding: 16px;
            border-radius: 8px;
            margin: 8px 0;
            line-height: 1.6;
            white-space: pre-wrap;
        ">
            {st.session_state.transcription.strip()}
        </div>
        """, unsafe_allow_html=True)
    
    # Translation section if enabled
    if st.session_state.settings["translate"]:
        with st.expander(f"🌍 Translation ({st.session_state.settings['target_lang'].upper()})", expanded=False):
            if st.session_state.translation is None:
                with st.spinner("Translating content..."):
                    st.session_state.translation = translate_text(
                        st.session_state.transcription,
                        st.session_state.settings["target_lang"]
                    )
            st.markdown(f"""
            <div style="
                background-color: rgba(0, 0, 0, 0.05);
                padding: 16px;
                border-radius: 8px;
                margin: 8px 0;
                line-height: 1.6;
                white-space: pre-wrap;
            ">
                {st.session_state.translation}
            </div>
            """, unsafe_allow_html=True)
    
    # Export buttons
    st.markdown("---")
    st.subheader("💾 Export Results")
    
    col1, col2 = st.columns(2)
    
    with col1:
        # Word document export
        doc = create_word_document(
            st.session_state.transcription,
            st.session_state.summary,
            st.session_state.translation if st.session_state.settings["translate"] else None
        )
        doc_path = os.path.join(tempfile.gettempdir(), f"transcription_{int(time.time())}.docx")
        doc.save(doc_path)
        
        with open(doc_path, "rb") as f:
            doc_bytes = f.read()
        
        st.download_button(
            label="📄 Download as Word Document",
            data=doc_bytes,
            file_name=f"transcription_report_{time.strftime('%Y%m%d')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )
    
    with col2:
        # PDF export (placeholder - would require additional libraries like pdfkit)
        st.button(
            "📝 Export as PDF (Coming Soon)",
            disabled=True,
            help="PDF export feature coming in next version",
            use_container_width=True
        )

# Main app function
def main():
    """Main application function"""
    # Custom CSS for better styling
    st.markdown("""
    <style>
        .stApp {
            max-width: 1200px;
            margin: 0 auto;
        }
        .stButton>button {
            transition: all 0.3s ease;
        }
        .stButton>button:hover {
            transform: scale(1.02);
        }
        .stExpander {
            border: 1px solid rgba(0, 0, 0, 0.1);
            border-radius: 8px;
            padding: 12px;
        }
        .stProgress>div>div>div {
            background-color: #4CAF50;
        }
    </style>
    """, unsafe_allow_html=True)
    
    st.title("🧠 AI Voice Transcriber & Reporter")
    st.markdown("""
    <div style="
        background-color: rgba(0, 0, 0, 0.05);
        padding: 16px;
        border-radius: 8px;
        margin-bottom: 24px;
    ">
        Upload an audio or video file to generate a professional transcription report with AI-powered summary.
    </div>
    """, unsafe_allow_html=True)
    
    # Setup sidebar and get settings
    st.session_state.settings = setup_sidebar()
    
    # File upload section
    uploaded_file = st.file_uploader(
        "Drag and drop or select an audio/video file",
        type=SUPPORTED_AUDIO_EXTENSIONS + SUPPORTED_VIDEO_EXTENSIONS,
        accept_multiple_files=False,
        key="file_uploader"
    )
    
    if uploaded_file is not None:
        # Check file size
        if uploaded_file.size > MAX_FILE_SIZE_MB * 1024 * 1024:
            st.error(f"❌ File size exceeds maximum limit of {MAX_FILE_SIZE_MB}MB")
            return
        
        # Save uploaded file to temp location
        temp_dir = tempfile.gettempdir()
        file_path = os.path.join(temp_dir, uploaded_file.name)
        
        with open(file_path, "wb") as f:
            f.write(uploaded_file.getbuffer())
        
        # Display file info
        display_file_info(file_path)
        
        # Check if file needs conversion
        file_ext = os.path.splitext(file_path)[1].lower()
        if file_ext not in [".wav"]:
            with st.spinner("🔄 Converting file to compatible format..."):
                converted_file = convert_to_wav(file_path)
                if not converted_file:
                    st.error("❌ File conversion failed. Please try another file.")
                    return
                file_path = converted_file
        
        # Process file when button is clicked
        if st.button("🚀 Start Processing", type="primary"):
            # Reset session state for new processing
            st.session_state.transcription = ""
            st.session_state.summary = None
            st.session_state.translation = None
            st.session_state.transcription_chunks = []
            st.session_state.translation_chunks = []
            st.session_state.processing_complete = False
            st.session_state.current_chunk = 0
            
            # Split audio into chunks for live processing
            with st.spinner("🔪 Splitting audio into chunks for live processing..."):
                st.session_state.transcription_chunks = split_audio(file_path)
            
            if not st.session_state.transcription_chunks:
                st.error("❌ Failed to split audio file. Please try again.")
                return
            
            # Start live processing if enabled
            if st.session_state.settings["live_mode"]:
                display_live_transcription()
            else:
                # Process all at once
                with st.spinner("🔊 Transcribing full audio (this may take a while)..."):
                    model = load_whisper_model(st.session_state.settings["whisper_model"])
                    result = model.transcribe(file_path)
                    st.session_state.transcription = result["text"]
                    st.session_state.processing_complete = True
                    st.success("✅ Transcription complete!")
            
            # Display final results when processing is complete
            if st.session_state.processing_complete:
                display_final_results()
    
    # Display previous results if available
    elif st.session_state.transcription and st.session_state.processing_complete:
        st.info("ℹ️ Showing results from previous session")
        display_final_results()

if __name__ == "__main__":
    # Check if CUDA is available
    if not torch.cuda.is_available():
        st.warning("⚠️ CUDA is not available. Processing might be slower on CPU.")
    
=======
import os
import time
import tempfile
import numpy as np
import streamlit as st
from pydub import AudioSegment
import whisper
from transformers import pipeline
from nltk.tokenize import sent_tokenize
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import torch
import ffmpeg
from typing import Optional, Tuple, List
from streamlit.components.v1 import html
import base64

# Set up NLTK
try:
    sent_tokenize("Test sentence.")
except:
    import nltk
    nltk.download('punkt')

# App configuration
st.set_page_config(
    page_title="🧠 AI Voice Transcriber & Reporter",
    page_icon="🎙️",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Constants
SUPPORTED_AUDIO_EXTENSIONS = [".mp3", ".wav", ".m4a", ".ogg", ".aac"]
SUPPORTED_VIDEO_EXTENSIONS = [".mp4", ".mov", ".avi"]
SUMMARIZATION_MODEL = "facebook/bart-large-cnn"
WHISPER_MODEL = "base"  # Default model
MAX_FILE_SIZE_MB = 100
CHUNK_SIZE_SECONDS = 20  # For live transcription

# Session state initialization
if 'transcription' not in st.session_state:
    st.session_state.transcription = None
if 'summary' not in st.session_state:
    st.session_state.summary = None
if 'translation' not in st.session_state:
    st.session_state.translation = None
if 'whisper_model' not in st.session_state:
    st.session_state.whisper_model = None
if 'summarizer' not in st.session_state:
    st.session_state.summarizer = None
if 'translation_pipeline' not in st.session_state:
    st.session_state.translation_pipeline = None
if 'transcription_chunks' not in st.session_state:
    st.session_state.transcription_chunks = []
if 'translation_chunks' not in st.session_state:
    st.session_state.translation_chunks = []
if 'processing_complete' not in st.session_state:
    st.session_state.processing_complete = False
if 'current_chunk' not in st.session_state:
    st.session_state.current_chunk = 0

# Helper functions
def convert_to_wav(input_file: str) -> str:
    """Convert any audio file to WAV format for processing"""
    output_file = os.path.join(tempfile.gettempdir(), f"converted_{int(time.time())}.wav")
    
    try:
        audio = AudioSegment.from_file(input_file)
        audio.export(output_file, format="wav")
    except Exception as e:
        st.error(f"Conversion failed with pydub: {str(e)}. Trying ffmpeg...")
        try:
            (
                ffmpeg
                .input(input_file)
                .output(output_file, ac=1, ar=16000)
                .run(quiet=True, overwrite_output=True)
            )
        except Exception as e:
            st.error(f"FFmpeg conversion also failed: {str(e)}")
            return None
    
    return output_file

def split_audio(file_path: str, chunk_size: int = CHUNK_SIZE_SECONDS) -> List[str]:
    """Split audio into chunks for live processing"""
    try:
        audio = AudioSegment.from_file(file_path)
        duration_ms = len(audio)
        chunk_ms = chunk_size * 1000
        chunks = []
        
        for i in range(0, duration_ms, chunk_ms):
            chunk = audio[i:i+chunk_ms]
            chunk_path = os.path.join(tempfile.gettempdir(), f"chunk_{i//1000}.wav")
            chunk.export(chunk_path, format="wav")
            chunks.append(chunk_path)
        
        return chunks
    except Exception as e:
        st.error(f"Failed to split audio: {str(e)}")
        return []

def load_whisper_model(model_name: str = "base") -> whisper.Whisper:
    """Load Whisper model with caching"""
    if st.session_state.whisper_model is None or st.session_state.whisper_model.name != model_name:
        with st.spinner(f"Loading Whisper {model_name} model (first time may take a while)..."):
            st.session_state.whisper_model = whisper.load_model(model_name)
    return st.session_state.whisper_model

def load_summarizer() -> pipeline:
    """Load summarization pipeline with caching"""
    if st.session_state.summarizer is None:
        with st.spinner("Loading summarization model..."):
            st.session_state.summarizer = pipeline("summarization", model=SUMMARIZATION_MODEL)
    return st.session_state.summarizer

def load_translation_pipeline(target_lang: str = "en") -> Optional[pipeline]:
    """Load translation pipeline for target language"""
    model_map = {
        "en": "Helsinki-NLP/opus-mt-mul-en",
        "it": "Helsinki-NLP/opus-mt-en-it",
        "fr": "Helsinki-NLP/opus-mt-en-fr",
        "es": "Helsinki-NLP/opus-mt-en-es",
        "de": "Helsinki-NLP/opus-mt-en-de"
    }
    
    if target_lang not in model_map:
        return None
    
    try:
        if st.session_state.translation_pipeline is None or st.session_state.translation_pipeline.model.config.lang2 != target_lang:
            with st.spinner(f"Loading translation model for {target_lang}..."):
                st.session_state.translation_pipeline = pipeline(
                    "translation", 
                    model=model_map[target_lang]
                )
        return st.session_state.translation_pipeline
    except Exception as e:
        st.error(f"Failed to load translation model: {str(e)}")
        return None

def transcribe_chunk(chunk_path: str, model: whisper.Whisper) -> str:
    """Transcribe a single audio chunk"""
    try:
        result = model.transcribe(chunk_path)
        return result["text"]
    except Exception as e:
        st.error(f"Error transcribing chunk: {str(e)}")
        return ""

def summarize_text(text: str, summary_length: str = "medium") -> str:
    """Summarize text using BART model"""
    summarizer = load_summarizer()
    
    length_params = {
        "short": {"min_length": 30, "max_length": 80},
        "medium": {"min_length": 60, "max_length": 150},
        "long": {"min_length": 120, "max_length": 300}
    }
    
    params = length_params.get(summary_length, length_params["medium"])
    
    sentences = sent_tokenize(text)
    chunks = []
    current_chunk = ""
    
    for sentence in sentences:
        if len(current_chunk) + len(sentence) < 1000:
            current_chunk += " " + sentence
        else:
            chunks.append(current_chunk)
            current_chunk = sentence
    
    if current_chunk:
        chunks.append(current_chunk)
    
    summaries = []
    progress_bar = st.progress(0)
    for i, chunk in enumerate(chunks):
        summary = summarizer(chunk, **params, do_sample=False)
        summaries.append(summary[0]['summary_text'])
        progress_bar.progress((i + 1) / len(chunks))
    
    return " ".join(summaries)

def translate_text(text: str, target_lang: str) -> str:
    """Translate text to target language"""
    translator = load_translation_pipeline(target_lang)
    if not translator:
        return None
    
    sentences = sent_tokenize(text)
    chunks = []
    current_chunk = ""
    
    for sentence in sentences:
        if len(current_chunk) + len(sentence) < 1000:
            current_chunk += " " + sentence
        else:
            chunks.append(current_chunk)
            current_chunk = sentence
    
    if current_chunk:
        chunks.append(current_chunk)
    
    translations = []
    progress_bar = st.progress(0)
    for i, chunk in enumerate(chunks):
        translation = translator(chunk)
        translations.append(translation[0]['translation_text'])
        progress_bar.progress((i + 1) / len(chunks))
    
    return " ".join(translations)

def create_word_document(transcription: str, summary: str, translation: Optional[str] = None) -> Document:
    """Create a professional Word document with the results"""
    doc = Document()
    
    # Set document margins
    for section in doc.sections:
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
    
    # Add title
    title = doc.add_paragraph()
    title_run = title.add_run("AI Voice Transcription Report")
    title_run.bold = True
    title_run.font.size = Pt(16)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Add subtitle with date
    doc.add_paragraph(time.strftime("%B %d, %Y"), style='Intense Quote')
    doc.add_paragraph("\n")
    
    # Add summary section
    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph(summary)
    doc.add_paragraph("\n")
    
    # Add transcription section
    doc.add_heading("Full Transcription", level=1)
    doc.add_paragraph(transcription)
    
    # Add translation if available
    if translation:
        doc.add_paragraph("\n")
        doc.add_heading("Translation", level=1)
        doc.add_paragraph(translation)
    
    # Add footer
    doc.add_paragraph("\n")
    footer = doc.add_paragraph()
    footer_run = footer.add_run("Generated by AI Voice Transcriber & Reporter")
    footer_run.italic = True
    footer_run.font.size = Pt(10)
    footer.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    return doc

def get_file_duration(file_path: str) -> float:
    """Get duration of audio/video file in seconds"""
    try:
        probe = ffmpeg.probe(file_path)
        duration = float(probe['format']['duration'])
        return duration
    except:
        return 0

def animate_text(text: str, speed: int = 5) -> None:
    """Create typing animation effect for text"""
    placeholder = st.empty()
    displayed_text = ""
    for char in text:
        displayed_text += char
        placeholder.markdown(f"""
        <div style="
            background-color: rgba(0, 0, 0, 0.05);
            padding: 12px;
            border-radius: 8px;
            margin: 8px 0;
            line-height: 1.6;
        ">
            {displayed_text}
        </div>
        """, unsafe_allow_html=True)
        time.sleep(1/speed)
    return displayed_text

# UI Components
def setup_sidebar():
    """Configure the sidebar options"""
    with st.sidebar:
        st.title("⚙️ Settings")
        
        # Model selection
        st.subheader("Model Options")
        whisper_model = st.selectbox(
            "Whisper Model Size",
            options=["tiny", "base", "small", "medium", "large"],
            index=1,
            help="Larger models are more accurate but require more memory and time"
        )
        
        summary_length = st.selectbox(
            "Summary Detail Level",
            options=["Brief", "Medium", "Detailed"],
            index=1,
            help="Control the length and detail of the generated summary"
        )
        
        # Translation options
        st.subheader("🌍 Translation")
        translate = st.checkbox("Enable Translation", False)
        target_lang = st.selectbox(
            "Target Language",
            options=["🇬🇧 English", "🇮🇹 Italian", "🇫🇷 French", "🇪🇸 Spanish", "🇩🇪 German"],
            index=0,
            disabled=not translate
        )
        
        # Theme options
        st.subheader("🎨 UI Preferences")
        theme = st.selectbox(
            "Color Theme",
            options=["Auto (System)", "Light", "Dark"],
            index=0
        )
        
        # Processing options
        st.subheader("⚡ Processing")
        live_mode = st.checkbox(
            "Live Transcription Mode", 
            True,
            help="Show transcription as it's being generated"
        )
        
        st.markdown("---")
        st.markdown("""
        **About this app:**
        - Powered by OpenAI Whisper for transcription
        - Uses Facebook BART for summarization
        - Helsinki-NLP for translation
        - Works fully offline after initial setup
        """)
    
    return {
        "whisper_model": whisper_model.lower(),
        "summary_length": summary_length.lower(),
        "translate": translate,
        "target_lang": target_lang[3:5].lower() if translate else None,
        "live_mode": live_mode,
        "theme": theme
    }

def display_file_info(file_path: str):
    """Display metadata about the uploaded file"""
    file_name = os.path.basename(file_path)
    file_size = os.path.getsize(file_path) / (1024 * 1024)  # in MB
    duration = get_file_duration(file_path)
    
    st.subheader("📄 File Information")
    col1, col2, col3 = st.columns(3)
    with col1:
        st.metric("File Name", file_name)
    with col2:
        st.metric("File Size", f"{file_size:.2f} MB")
    with col3:
        st.metric("Duration", f"{duration:.2f} sec" if duration > 0 else "Unknown")

def display_live_transcription():
    """Display transcription chunks as they're processed"""
    st.subheader("🔴 Live Transcription Progress")
    progress_bar = st.progress(0)
    status_text = st.empty()
    
    total_chunks = len(st.session_state.transcription_chunks)
    model = load_whisper_model(st.session_state.settings["whisper_model"])
    
    for i in range(st.session_state.current_chunk, total_chunks):
        chunk_path = st.session_state.transcription_chunks[i]
        
        # Update progress
        progress = (i + 1) / total_chunks
        progress_bar.progress(progress)
        status_text.markdown(f"""
        <div style="background-color: rgba(0, 0, 0, 0.05); padding: 12px; border-radius: 8px; margin: 8px 0;">
            Processing chunk {i+1} of {total_chunks}...
        </div>
        """, unsafe_allow_html=True)
        
        # Transcribe chunk
        chunk_text = transcribe_chunk(chunk_path, model)
        
        # Display chunk with animation
        st.markdown(f"""
        <div style="
            background-color: rgba(0, 0, 0, 0.05);
            padding: 12px;
            border-radius: 8px;
            margin: 8px 0;
            border-left: 4px solid #4CAF50;
        ">
            <strong>Chunk {i+1}:</strong><br>
        </div>
        """, unsafe_allow_html=True)
        
        animate_text(chunk_text)
        
        # Update full transcription
        if st.session_state.transcription is None:
            st.session_state.transcription = ""
        st.session_state.transcription += " " + chunk_text
        st.session_state.current_chunk = i + 1
    
    # Complete processing
    progress_bar.empty()
    status_text.success("✅ Transcription complete!")
    st.session_state.processing_complete = True

def display_final_results():
    """Display the final transcription, summary, and translation"""
    st.subheader("📝 Final Results")
    
    # Summary section
    with st.expander("📋 Executive Summary", expanded=True):
        if st.session_state.summary is None:
            with st.spinner("Generating summary..."):
                st.session_state.summary = summarize_text(
                    st.session_state.transcription,
                    st.session_state.settings["summary_length"]
                )
        st.markdown(f"""
        <div style="
            background-color: rgba(0, 0, 0, 0.05);
            padding: 16px;
            border-radius: 8px;
            margin: 8px 0;
            line-height: 1.6;
        ">
            {st.session_state.summary}
        </div>
        """, unsafe_allow_html=True)
    
    # Transcription section
    with st.expander("🎧 Full Transcription", expanded=False):
        st.markdown(f"""
        <div style="
            background-color: rgba(0, 0, 0, 0.05);
            padding: 16px;
            border-radius: 8px;
            margin: 8px 0;
            line-height: 1.6;
            white-space: pre-wrap;
        ">
            {st.session_state.transcription.strip()}
        </div>
        """, unsafe_allow_html=True)
    
    # Translation section if enabled
    if st.session_state.settings["translate"]:
        with st.expander(f"🌍 Translation ({st.session_state.settings['target_lang'].upper()})", expanded=False):
            if st.session_state.translation is None:
                with st.spinner("Translating content..."):
                    st.session_state.translation = translate_text(
                        st.session_state.transcription,
                        st.session_state.settings["target_lang"]
                    )
            st.markdown(f"""
            <div style="
                background-color: rgba(0, 0, 0, 0.05);
                padding: 16px;
                border-radius: 8px;
                margin: 8px 0;
                line-height: 1.6;
                white-space: pre-wrap;
            ">
                {st.session_state.translation}
            </div>
            """, unsafe_allow_html=True)
    
    # Export buttons
    st.markdown("---")
    st.subheader("💾 Export Results")
    
    col1, col2 = st.columns(2)
    
    with col1:
        # Word document export
        doc = create_word_document(
            st.session_state.transcription,
            st.session_state.summary,
            st.session_state.translation if st.session_state.settings["translate"] else None
        )
        doc_path = os.path.join(tempfile.gettempdir(), f"transcription_{int(time.time())}.docx")
        doc.save(doc_path)
        
        with open(doc_path, "rb") as f:
            doc_bytes = f.read()
        
        st.download_button(
            label="📄 Download as Word Document",
            data=doc_bytes,
            file_name=f"transcription_report_{time.strftime('%Y%m%d')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )
    
    with col2:
        # PDF export (placeholder - would require additional libraries like pdfkit)
        st.button(
            "📝 Export as PDF (Coming Soon)",
            disabled=True,
            help="PDF export feature coming in next version",
            use_container_width=True
        )

# Main app function
def main():
    """Main application function"""
    # Custom CSS for better styling
    st.markdown("""
    <style>
        .stApp {
            max-width: 1200px;
            margin: 0 auto;
        }
        .stButton>button {
            transition: all 0.3s ease;
        }
        .stButton>button:hover {
            transform: scale(1.02);
        }
        .stExpander {
            border: 1px solid rgba(0, 0, 0, 0.1);
            border-radius: 8px;
            padding: 12px;
        }
        .stProgress>div>div>div {
            background-color: #4CAF50;
        }
    </style>
    """, unsafe_allow_html=True)
    
    st.title("🧠 AI Voice Transcriber & Reporter")
    st.markdown("""
    <div style="
        background-color: rgba(0, 0, 0, 0.05);
        padding: 16px;
        border-radius: 8px;
        margin-bottom: 24px;
    ">
        Upload an audio or video file to generate a professional transcription report with AI-powered summary.
    </div>
    """, unsafe_allow_html=True)
    
    # Setup sidebar and get settings
    st.session_state.settings = setup_sidebar()
    
    # File upload section
    uploaded_file = st.file_uploader(
        "Drag and drop or select an audio/video file",
        type=SUPPORTED_AUDIO_EXTENSIONS + SUPPORTED_VIDEO_EXTENSIONS,
        accept_multiple_files=False,
        key="file_uploader"
    )
    
    if uploaded_file is not None:
        # Check file size
        if uploaded_file.size > MAX_FILE_SIZE_MB * 1024 * 1024:
            st.error(f"❌ File size exceeds maximum limit of {MAX_FILE_SIZE_MB}MB")
            return
        
        # Save uploaded file to temp location
        temp_dir = tempfile.gettempdir()
        file_path = os.path.join(temp_dir, uploaded_file.name)
        
        with open(file_path, "wb") as f:
            f.write(uploaded_file.getbuffer())
        
        # Display file info
        display_file_info(file_path)
        
        # Check if file needs conversion
        file_ext = os.path.splitext(file_path)[1].lower()
        if file_ext not in [".wav"]:
            with st.spinner("🔄 Converting file to compatible format..."):
                converted_file = convert_to_wav(file_path)
                if not converted_file:
                    st.error("❌ File conversion failed. Please try another file.")
                    return
                file_path = converted_file
        
        # Process file when button is clicked
        if st.button("🚀 Start Processing", type="primary"):
            # Reset session state for new processing
            st.session_state.transcription = ""
            st.session_state.summary = None
            st.session_state.translation = None
            st.session_state.transcription_chunks = []
            st.session_state.translation_chunks = []
            st.session_state.processing_complete = False
            st.session_state.current_chunk = 0
            
            # Split audio into chunks for live processing
            with st.spinner("🔪 Splitting audio into chunks for live processing..."):
                st.session_state.transcription_chunks = split_audio(file_path)
            
            if not st.session_state.transcription_chunks:
                st.error("❌ Failed to split audio file. Please try again.")
                return
            
            # Start live processing if enabled
            if st.session_state.settings["live_mode"]:
                display_live_transcription()
            else:
                # Process all at once
                with st.spinner("🔊 Transcribing full audio (this may take a while)..."):
                    model = load_whisper_model(st.session_state.settings["whisper_model"])
                    result = model.transcribe(file_path)
                    st.session_state.transcription = result["text"]
                    st.session_state.processing_complete = True
                    st.success("✅ Transcription complete!")
            
            # Display final results when processing is complete
            if st.session_state.processing_complete:
                display_final_results()
    
    # Display previous results if available
    elif st.session_state.transcription and st.session_state.processing_complete:
        st.info("ℹ️ Showing results from previous session")
        display_final_results()

if __name__ == "__main__":
    # Check if CUDA is available
    if not torch.cuda.is_available():
        st.warning("⚠️ CUDA is not available. Processing might be slower on CPU.")
    
>>>>>>> 9355cb45 (chore: initial commit)
    main()